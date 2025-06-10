import pandas as pd
import streamlit as st
from collections import defaultdict
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Ekintzen antolatzailea", layout="wide")
st.title("📋 EKE-en antolatzailea")

uploaded_file = st.file_uploader("📂 Igo formularioaren CSV fitxategia", type="csv")

if uploaded_file:
    df = pd.read_csv(uploaded_file)

    alumno_col = df.columns[2]
    curso_col = df.columns[4]
    actividad_cols = [col for col in df.columns if "ekintzak" in col.lower() or "actividad" in col.lower()]

    actividades_dict = defaultdict(lambda: defaultdict(list))

    for _, row in df.iterrows():
        alumno = str(row[alumno_col]).strip().title()
        curso = row[curso_col]
        for col in actividad_cols:
            actividad = row[col]
            if pd.notna(actividad):
                actividades = [a.strip() for a in str(actividad).split(",")]
                for act in actividades:
                    actividades_dict[act][curso].append(alumno)

    # Crear documento .docx antes de mostrar en pantalla
    doc = Document()
    doc.add_heading("Ekintzen zerrenda", 0)

    for actividad in sorted(actividades_dict):
        doc.add_heading(actividad, level=1)
        for curso in sorted(actividades_dict[actividad]):
            doc.add_paragraph(f"📘 {curso}", style='Heading 2')
            for alumno in sorted(actividades_dict[actividad][curso]):
                doc.add_paragraph(alumno, style='List Bullet')  # <- Sin el guion

    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Mostrar el botón al principio
    st.download_button(
        label="⬇️ Deskargatu DOCX fitxategia",
        data=docx_buffer,
        file_name="ekintzak_lista.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Mostrar en pantalla
    for actividad in sorted(actividades_dict):
        st.header(f"🟢 {actividad}")
        for curso in sorted(actividades_dict[actividad]):
            st.subheader(f"📘 {curso}")
            for alumno in sorted(actividades_dict[actividad][curso]):
                st.write(f"- {alumno}")

else:
    st.info("Mesedez, igo CSV fitxategi bat hasteko.")
